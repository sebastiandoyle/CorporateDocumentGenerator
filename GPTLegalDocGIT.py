import pdfplumber
import datetime
import csv
import pandas as pd
import numpy as np
import pickle
import os
import tkinter as tk
from tkinter import filedialog
from openai.embeddings_utils import get_embedding, cosine_similarity
import openai
import os
import nltk
import requests
import json
import tkinter as tk
from tkinter import filedialog
from csv import DictWriter
from tkinter import messagebox
from pandas import read_csv
from numpy import fromstringW
import tkinter as tk
from tkinter import filedialog, messagebox
import concurrent.futures
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

class DocWriter:
    def __init__(self, title):
        self.doc = Document()
        self.doc.add_heading(title, level=0)

    def add_section(self, heading, content):
        self.doc.add_heading(heading, level=1)
        self.doc.add_paragraph(content)

    def save(self, filename):
        self.doc.save(filename)

class App:
    def __init__(self, root):
        self.root = root
        self.create_widgets()
        self.pdf_folder_paths = []
        self.output_folder = ""
        self.csv_output_path = ""
        
    def create_widgets(self):
        self.label1 = tk.Label(self.root, text="PDF Folder Paths")
        self.label1.pack()
        self.pdf_paths_listbox = tk.Listbox(self.root)
        self.pdf_paths_listbox.pack()
        self.add_pdf_path_button = tk.Button(self.root, text="Add PDF Path", command=self.add_pdf_path)
        self.add_pdf_path_button.pack()
        
        self.label2 = tk.Label(self.root, text="Output Folder")
        self.label2.pack()
        self.output_folder_entry = tk.Entry(self.root)
        self.output_folder_entry.pack()
        self.browse_button = tk.Button(self.root, text="Browse", command=self.browse_output_folder)
        self.browse_button.pack()
        
        self.search_label = tk.Label(self.root, text="Search Phrase:")
        self.search_label.pack()
        self.search_entry = tk.Entry(self.root)
        self.search_entry.pack()
        self.search_button = tk.Button(self.root, text="Search", command=self.search)
        self.search_button.pack()
        
    def add_pdf_path(self):
        pdf_path = filedialog.askdirectory(title="Select PDF Folder")
        if pdf_path:
            self.pdf_folder_paths.append(pdf_path)
            self.pdf_paths_listbox.insert(tk.END, pdf_path)
        
    def browse_output_folder(self):
        output_folder = filedialog.askdirectory(title='Select output directory')
        if output_folder:
            today_date = datetime.datetime.now().strftime('%Y-%m-%d')
            self.output_folder_entry.insert(0, output_folder)
            self.output_folder = output_folder
            self.csv_output_path = os.path.join(output_folder, f"{today_date}_combined_csv.csv")
    
    def search(self):
        search_phrase = self.search_entry.get()
        if not search_phrase:
            messagebox.showwarning("Search Phrase Required", "Please enter a search phrase.")
            return
        
        if not self.output_folder:
            messagebox.showwarning("Output Folder Required", "Please select an output folder.")
            return
        results = self.run_search(self.pdf_folder_paths, self.output_folder, search_phrase)
        print("Search results:", results)

    def run_search(self, pdf_folder_paths, output_folder, search_phrase):
            
        nltk.download('punkt')

        openai.api_key = "YOUR API KEY"

        API_KEY = "YOUR API KEY"
        API_URL = "https://api.openai.com/v1/chat/completions"
        MODEL_NAME = "gpt-4"

        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {API_KEY}",
        }

        def send_gpt_request(gpt_command, combined_results, search_phrase):
            prompt = [
                {"role": "system", "content": f"{gpt_command}"},
                {"role": "user", "content": f"This is the user's request: \"{search_phrase}\". Here is some context from some documents, it may or may not be relevant: \"{combined_results}\""}
            ]


            request_body = {
                "model": MODEL_NAME,
                "messages": prompt
            }

            response = requests.post(API_URL, headers=headers, json=request_body)

            if response.status_code == 200:
                return response.json()["choices"][0]["message"]["content"].strip()
            else:
                if response.status_code == 400:
                    print(f"Error 400 response: {response.text}")
                raise Exception(f"Error in GPT-4 API request: {response.text}")

        def load_embeddings(pickle_path):
            if os.path.isfile(pickle_path):
                with open(pickle_path, 'rb') as f:
                    return pickle.load(f)
            else:
                return {}

        def save_embeddings(embeddings, pickle_path):
            with open(pickle_path, 'wb') as f:
                pickle.dump(embeddings, f)

        def read_pdf_and_save_to_csv(pdf_file, output_csv, embeddings_pickle_path, df):
            with pdfplumber.open(pdf_file) as pdf:
                text = ""

                for page in pdf.pages:
                    text += page.extract_text()

                sentences = nltk.sent_tokenize(text)
                target_word_count = 200
                chunks = []
                current_chunk = []
                current_word_count = 0

                for sentence in sentences:
                    words = sentence.split()
                    word_count = len(words)
                    if current_word_count + word_count <= target_word_count:
                        current_chunk.append(sentence)
                        current_word_count += word_count
                    else:
                        chunks.append(' '.join(current_chunk))
                        current_chunk = [sentence]
                        current_word_count = word_count

                # Add the last chunk
                if current_chunk:
                    chunks.append(' '.join(current_chunk))

                embeddings = load_embeddings(embeddings_pickle_path)

                with open(output_csv, 'a', newline='', encoding='utf-8') as csvfile:
                    writer = csv.writer(csvfile)
                    for chunk in chunks:
                        # Calculate and save the embedding only if the chunk is not empty and not in the DataFrame already
                        if chunk and chunk not in df['text'].values:
                            embedding = get_embedding(chunk, engine="text-embedding-ada-002")
                            embeddings[chunk] = embedding
                            writer.writerow([chunk, pdf_file])
                            
                # Save the updated embeddings
                save_embeddings(embeddings, embeddings_pickle_path)

        def search_text_chunks(df, search_body, embeddings_pickle_path, n=10, pprint=True):
            embeddings = load_embeddings(embeddings_pickle_path)

            phrase_embedding = get_embedding(
                search_body,
                engine="text-embedding-ada-002"
            )
            
            df["similarity"] = df.text.apply(lambda x: cosine_similarity(embeddings[x], phrase_embedding) if x in embeddings else 0)

            results = (
                df.sort_values("similarity", ascending=False)
                .head(n)
                .text
            )

            # Remove new lines by replacing them with spaces
            cleaned_results = [result.replace('\n', ' ') for result in results]

            if pprint:
                for r in cleaned_results:
                    print(r[:200])
                    print()

            return results
        
        def generate_document_section(section, combined_results, search_phrase):
            # Pass combined search results and user query to GPT-4 for filling out a section
            gpt_command = f"Provide detailed information on \"{section}\""
            
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(send_gpt_request, gpt_command, combined_results, search_phrase)
                return future.result()


        today_date = datetime.datetime.now().strftime('%Y-%m-%d')
        csv_output_path = os.path.join(output_folder, f"{today_date}_combined_csv.csv")

        print("Creating CSV file with header...")

        # Create the CSV file with a header
        with open(csv_output_path, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            writer.writerow(['text', 'file'])

        embeddings_pickle_path = os.path.join(output_folder, "embeddings.pickle")

        # Load the CSV file into a DataFrame
        if os.path.isfile(csv_output_path):
            df = pd.read_csv(csv_output_path)
        else:
            df = pd.DataFrame(columns=['text', 'file'])

        with concurrent.futures.ThreadPoolExecutor(max_workers=11) as executor:
            # First, create all futures
            future_to_pdf = {executor.submit(read_pdf_and_save_to_csv, os.path.join(dirpath, pdf_file), csv_output_path, embeddings_pickle_path, df): pdf_file
                            for folder_path in pdf_folder_paths
                            for dirpath, dirnames, filenames in os.walk(folder_path)
                            for pdf_file in filenames
                            if pdf_file.endswith('.pdf')}
            
            # Then, as they complete, process them
            for future in concurrent.futures.as_completed(future_to_pdf):
                pdf_file = future_to_pdf[future]
                try:
                    future.result()
                    print(f"Processing {pdf_file}... Done!")
                except Exception as exc:
                    print(f"Processing {pdf_file}... ERROR: {exc}")

        print("Reading CSV file and creating embeddings...")
        df = pd.read_csv(csv_output_path)


        # Continuously prompt user for search phrases and execute search
        while True:
            while True:
                search_phrase = input("Enter the search phrase (or type 'exit' to quit): ")
                if search_phrase.lower() == 'exit':
                    break
                print("Searching for the phrase...")
                search_body = search_phrase
                results = search_text_chunks(df, search_body, embeddings_pickle_path, n=4)
                results = [result.replace('\n', ' ') for result in results]  # Remove new lines by replacing them with spaces
                print(results) # Just for debugging to see what context it is using

                # Pass combined search results and user query to GPT-4 for a recommended approach and summary of 1st batch of search results
                combined_results = ' '.join(results)
                gpt_command = "Based on the query, what should be the structure of the help document? VERY IMPORTANT: REPLY ONLY WITH HEADINGS SEPARATED BY COMMAS!!"
                document_structure = send_gpt_request(gpt_command, combined_results, search_phrase)
                print("Proposed Document Structure:")
                print(document_structure)

                sections = document_structure.split(',')

                document_content = []

                with concurrent.futures.ThreadPoolExecutor() as executor:
                    # Prepare all the futures
                    future_to_section = {executor.submit(generate_document_section, section, combined_results, search_phrase): section for section in sections}
                    # As each future completes, add its result to document_content
                    for future in concurrent.futures.as_completed(future_to_section):
                        section = future_to_section[future]
                        try:
                            section_content = future.result()
                            document_content.append(section_content)
                        except Exception as exc:
                            print(f"An error occurred while generating content for section {section}: {exc}")
                                                
                document_writer = DocWriter("Generated Document")

                for section, section_content in zip(sections, document_content):
                    document_writer.add_section(section, section_content)

                document_writer.save(os.path.join(output_folder, f"{today_date}_generated_doc.docx"))

            if input("To start a new search, press any key. To quit, type 'exit': ").lower() == 'exit':
                break

            
root = tk.Tk()
app = App(root)
root.mainloop()
